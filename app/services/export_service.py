from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.db.models import Report

def generate_word_report(report: Report):
    doc = Document()
    
    # Title
    title = doc.add_heading('Periodiek Systeembeheer Rapport', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Customer Info
    doc.add_heading('Klantgegevens', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = 'Klantnaam:'
    table.cell(0, 1).text = report.customer.name
    table.cell(1, 0).text = 'Locatie:'
    table.cell(1, 1).text = report.customer.location or 'N/A'
    table.cell(2, 0).text = 'Datum:'
    table.cell(2, 1).text = report.date_performed.strftime('%d-%m-%Y')
    
    doc.add_paragraph()
    
    # Results
    doc.add_heading('Checklist Resultaten', level=1)
    
    # Sort items by category order and then checkpoint name if possible
    # For now, just list them
    results_table = doc.add_table(rows=1, cols=3)
    results_table.style = 'Table Grid'
    hdr_cells = results_table.rows[0].cells
    hdr_cells[0].text = 'Checkpoint'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Opmerking'
    
    for item in report.items:
        row_cells = results_table.add_row().cells
        # We need to access the checkpoint name. 
        # ReportItem has checkpoint relationship
        # Let's ensure models are correct for this.
        checkpoint_name = "Unknown"
        if hasattr(item, 'checkpoint') and item.checkpoint:
            checkpoint_name = item.checkpoint.name
        elif item.checkpoint_id:
            # Fallback if relationship not loaded
            checkpoint_name = f"Checkpoint {item.checkpoint_id}"
            
        row_cells[0].text = checkpoint_name
        row_cells[1].text = str(item.result.value)
        row_cells[2].text = item.comment or ''
        
    return doc
